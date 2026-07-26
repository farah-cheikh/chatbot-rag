from docx import Document
def extract_docx(file_path):
    doc=Document(file_path)

    text= []
from pathlib import Path
from docx import Document

def extract_docx(file_path):
    doc = Document(file_path)
    text = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    cells.append(value)
            if cells:
                text.append(" | ".join(cells))
    return "\n".join(text)


def load_documents(folder_path):
    folder = Path(folder_path)
    documents = []
    for file in folder.glob("*.docx"):
        text = extract_docx(file)
        documents.append({
            "filename": file.name,
            "text": text
        })
    return documents